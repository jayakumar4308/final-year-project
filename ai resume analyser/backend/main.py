from fastapi import (
    FastAPI,
    UploadFile,
    File,
    HTTPException
)

from fastapi.middleware.cors import CORSMiddleware

import pymupdf

from docx import Document

import io


# =========================================================
# IMPORT OUR ANALYZER AND JOB RECOMMENDER
# =========================================================

from analyzer import analyze_resume

from recommender import recommend_jobs


# =========================================================
# CREATE FASTAPI APP
# =========================================================

app = FastAPI(
    title="AI Resume Analyzer API",
    version="1.0.0"
)


# =========================================================
# CORS
# =========================================================

app.add_middleware(

    CORSMiddleware,

    allow_origins=["*"],

    allow_credentials=True,

    allow_methods=["*"],

    allow_headers=["*"]
)


# =========================================================
# PDF TEXT EXTRACTION
# =========================================================

def extract_pdf_text(file_bytes):

    text = ""

    try:

        pdf = pymupdf.open(
            stream=file_bytes,
            filetype="pdf"
        )


        for page in pdf:

            text += page.get_text()


        pdf.close()


    except Exception as e:

        raise HTTPException(

            status_code=400,

            detail=f"Could not read PDF file: {str(e)}"
        )


    return text


# =========================================================
# DOCX TEXT EXTRACTION
# =========================================================

def extract_docx_text(file_bytes):

    try:

        document = Document(
            io.BytesIO(file_bytes)
        )


        text = "\n".join(

            paragraph.text

            for paragraph in document.paragraphs

            if paragraph.text.strip()
        )


        # Read tables also

        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text.strip():

                        text += "\n" + cell.text


        return text


    except Exception as e:

        raise HTTPException(

            status_code=400,

            detail=f"Could not read DOCX file: {str(e)}"
        )


# =========================================================
# HOME
# =========================================================

@app.get("/")
def home():

    return {

        "message": "AI Resume Analyzer Backend is running!",

        "status": "active"
    }


# =========================================================
# ANALYZE RESUME
# =========================================================

@app.post("/analyze")
async def analyze_resume_api(

    file: UploadFile = File(...)
):


    # -----------------------------------------------------
    # Allowed file types
    # -----------------------------------------------------

    allowed_types = [

        "application/pdf",

        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    ]


    # -----------------------------------------------------
    # Validate file type
    # -----------------------------------------------------

    if file.content_type not in allowed_types:

        raise HTTPException(

            status_code=400,

            detail="Only PDF and DOCX files are supported."
        )


    # -----------------------------------------------------
    # Read uploaded file
    # -----------------------------------------------------

    file_bytes = await file.read()


    if not file_bytes:

        raise HTTPException(

            status_code=400,

            detail="Uploaded file is empty."
        )


    # -----------------------------------------------------
    # Extract text
    # -----------------------------------------------------

    if file.content_type == "application/pdf":

        resume_text = extract_pdf_text(
            file_bytes
        )

    else:

        resume_text = extract_docx_text(
            file_bytes
        )


    # -----------------------------------------------------
    # Check extracted text
    # -----------------------------------------------------

    if not resume_text.strip():

        raise HTTPException(

            status_code=400,

            detail=(
                "Could not extract text from the resume. "
                "Please upload a text-based PDF or DOCX file."
            )
        )


    # =====================================================
    # AI RESUME ANALYSIS
    # =====================================================

    analysis = analyze_resume(
        resume_text
    )


    # =====================================================
    # JOB RECOMMENDATION
    # =====================================================

    recommended_jobs = recommend_jobs(

        analysis["skills"]
    )


    # =====================================================
    # FINAL RESPONSE
    # =====================================================

    return {

        "message": "Resume analyzed successfully",

        "filename": file.filename,

        "score": analysis["score"],

        "skills": analysis["skills"],

        "sections": analysis["sections"],

        "suggestions": analysis["suggestions"],

        "recommended_jobs": recommended_jobs,

        "word_count": analysis["word_count"],

        "text": resume_text
    }